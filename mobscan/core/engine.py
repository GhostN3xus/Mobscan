"""
Test Engine - Core orchestration engine

Manages the execution of security tests, coordinates between modules,
aggregates findings, and produces reports.
"""

import asyncio
import logging
from typing import List, Dict, Optional, Any
from datetime import datetime
from pathlib import Path
import concurrent.futures
import json
import hashlib
import zipfile
import xml.etree.ElementTree as ET

from ..models.scan_result import ScanResult, ApplicationInfo, TestCoverage, RiskMetrics
from ..models.finding import Finding, Severity
from ..models.masvs_mapping import MAVSMapping, MAVSLevel
from .config import MobscanConfig

# Report generation imports
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


logger = logging.getLogger(__name__)


class TestEngine:
    """
    Main orchestration engine for mobile security testing.

    Responsibilities:
    - Load and validate configuration
    - Coordinate between test modules
    - Manage test execution workflow
    - Aggregate findings and results
    - Generate reports
    """

    def __init__(self, config: Optional[MobscanConfig] = None):
        """
        Initialize the test engine.

        Args:
            config: MobscanConfig instance. If None, uses default config.
        """
        self.config = config or MobscanConfig.default_config()
        self.logger = self._setup_logging()
        self.scan_result: Optional[ScanResult] = None
        self.test_modules: Dict[str, Any] = {}

    def _setup_logging(self) -> logging.Logger:
        """Setup logging configuration"""
        logger = logging.getLogger("mobscan.engine")
        logger.setLevel(getattr(logging, self.config.log_level))

        # Only add handler if logger doesn't already have handlers
        # This prevents duplicate handlers on multiple initializations
        if not logger.handlers:
            # Console handler
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)

        return logger

    def initialize_scan(self, app_path: str, app_name: str = "") -> ScanResult:
        """
        Initialize a new scan and create ScanResult object.

        Args:
            app_path: Path to the APK/IPA file
            app_name: Display name for the application

        Returns:
            ScanResult: Initialized scan result object
        """
        self.logger.info(f"Initializing scan for {app_path}")

        # Validate app file
        if not Path(app_path).exists():
            raise FileNotFoundError(f"Application file not found: {app_path}")

        # Create scan result
        self.scan_result = ScanResult()
        self.scan_result.scan_name = app_name or Path(app_path).stem
        self.scan_result.scan_intensity = self.config.scan_intensity.value
        self.scan_result.scan_modules = self.config.modules_enabled
        self.scan_result.assessment_type = "Automated"

        # Extract app information
        self.scan_result.app_info = self._extract_app_info(app_path)

        # Initialize test coverage
        self.scan_result.test_coverage = TestCoverage(
            modules_executed=self.config.modules_enabled
        )

        self.logger.info(f"Scan initialized: {self.scan_result.scan_id}")
        return self.scan_result

    def _extract_app_info(self, app_path: str) -> ApplicationInfo:
        """
        Extract application information from APK/IPA.
        """
        app_file = Path(app_path)

        if not app_file.exists():
            raise FileNotFoundError(f"Application file not found: {app_path}")

        file_size = app_file.stat().st_size
        platform = "android" if app_file.suffix.lower() == ".apk" else "ios"

        # Calculate SHA256 hash
        file_hash = self._calculate_file_hash(app_path)

        # Extract platform-specific info
        if platform == "android":
            package_name, version, app_name = self._extract_apk_info(app_path)
        else:
            # For iOS, extract basic info (real extraction would need plistlib)
            package_name = app_file.stem
            version = "1.0.0"
            app_name = app_file.stem

        return ApplicationInfo(
            app_name=app_name or app_file.stem,
            package_name=package_name or app_file.stem,
            version=version or "1.0.0",
            platform=platform,
            file_size=file_size,
            file_hash=file_hash,
        )

    def _calculate_file_hash(self, filepath: str) -> str:
        """Calculate SHA256 hash of a file"""
        sha256_hash = hashlib.sha256()
        try:
            with open(filepath, "rb") as f:
                for byte_block in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(byte_block)
            return sha256_hash.hexdigest()
        except Exception as e:
            self.logger.warning(f"Failed to calculate file hash: {e}")
            return ""

    def _extract_apk_info(self, apk_path: str) -> tuple[str, str, str]:
        """
        Extract package name, version, and app name from APK manifest.

        Returns:
            Tuple of (package_name, version, app_name)
        """
        try:
            with zipfile.ZipFile(apk_path, 'r') as apk:
                # Read AndroidManifest.xml
                manifest_data = apk.read('AndroidManifest.xml')

                # Parse binary XML (simplified approach)
                # For production, use androguard library
                package_name = "com.app"
                version = "1.0"
                app_name = Path(apk_path).stem

                # Try to extract from manifest if readable
                try:
                    manifest_text = manifest_data.decode('utf-8', errors='ignore')
                    if 'package=' in manifest_text:
                        # Simple extraction (real parsing would use XML parser)
                        parts = manifest_text.split('package=')
                        if len(parts) > 1:
                            package_name = parts[1].split('"')[1]
                    if 'versionName=' in manifest_text:
                        parts = manifest_text.split('versionName=')
                        if len(parts) > 1:
                            version = parts[1].split('"')[1]
                except:
                    pass

                return package_name, version, app_name
        except Exception as e:
            self.logger.warning(f"Failed to extract APK info: {e}")
            return Path(apk_path).stem, "1.0.0", Path(apk_path).stem

    def load_test_modules(self):
        """Load and register all test modules"""
        self.logger.info("Loading test modules...")

        # Module loading logic would go here
        # For now, this is a placeholder

        if self.config.is_module_enabled("sast"):
            self.logger.info("  - SAST module loaded")
            self.test_modules["sast"] = {}

        if self.config.is_module_enabled("dast"):
            self.logger.info("  - DAST module loaded")
            self.test_modules["dast"] = {}

        if self.config.is_module_enabled("frida"):
            self.logger.info("  - Frida module loaded")
            self.test_modules["frida"] = {}

        self.logger.info(f"Total modules loaded: {len(self.test_modules)}")

    def execute_tests(self) -> ScanResult:
        """
        Execute all configured tests.

        Returns:
            ScanResult: Completed scan result with all findings
        """
        if not self.scan_result:
            raise RuntimeError("Scan not initialized. Call initialize_scan first.")

        self.logger.info("Starting test execution...")
        self.load_test_modules()

        try:
            # Execute tests in parallel
            self._execute_parallel_tests()

            # Deduplicate findings
            duplicates = self.scan_result.deduplicate_findings()
            self.logger.info(f"Removed {duplicates} duplicate findings")

            # Map findings to MASVS
            self._map_findings_to_masvs()

            # Calculate final metrics
            self.scan_result.finalize()
            self.logger.info(f"Scan completed. Total findings: {len(self.scan_result.findings)}")

        except Exception as e:
            self.logger.error(f"Error during test execution: {str(e)}", exc_info=True)
            raise

        return self.scan_result

    def _execute_parallel_tests(self):
        """Execute tests in parallel using thread pool"""
        executor = concurrent.futures.ThreadPoolExecutor(
            max_workers=self.config.parallel_workers
        )

        futures = []

        # Submit SAST tests
        if "sast" in self.test_modules:
            self.logger.info("Submitting SAST tests...")
            futures.append(executor.submit(self._run_sast_tests))

        # Submit DAST tests
        if "dast" in self.test_modules:
            self.logger.info("Submitting DAST tests...")
            futures.append(executor.submit(self._run_dast_tests))

        # Submit Frida tests
        if "frida" in self.test_modules:
            self.logger.info("Submitting Frida instrumentation tests...")
            futures.append(executor.submit(self._run_frida_tests))

        # Wait for all tests to complete
        for future in concurrent.futures.as_completed(futures):
            try:
                future.result(timeout=self.config.timeout_global)
            except Exception as e:
                self.logger.error(f"Test execution failed: {str(e)}")

        executor.shutdown(wait=True)

    def _run_sast_tests(self):
        """Execute static analysis tests"""
        self.logger.info("Running SAST tests...")

        # Sample findings for demonstration
        # In production, integrate with MobSF, JADX, etc.

        finding = Finding(
            id="FINDING-001",
            title="Hardcoded API Credentials",
            description="API credentials found hardcoded in source code",
            severity=Severity.CRITICAL,
            cvss=self._create_cvss_score(9.8),
            cwe=["CWE-798"],
            owasp_category="A02:2021 - Cryptographic Failures",
            test_name="Hardcoded Secrets Detection",
            module="sast",
            mastg_category="MASTG-STORAGE-1",
            masvs_category="MSTG-STORAGE-1",
            affected_component="NetworkManager.java"
        )

        self.scan_result.add_finding(finding)
        self.logger.info("SAST tests completed")

    def _run_dast_tests(self):
        """Execute dynamic analysis tests"""
        self.logger.info("Running DAST tests...")

        # Sample findings
        finding = Finding(
            id="FINDING-002",
            title="Insecure Certificate Pinning",
            description="Certificate pinning not implemented or incorrectly implemented",
            severity=Severity.HIGH,
            cvss=self._create_cvss_score(7.5),
            cwe=["CWE-295"],
            owasp_category="A04:2021 - Insecure Design",
            test_name="Certificate Pinning Validation",
            module="dast",
            mastg_category="MASTG-NET-2",
            masvs_category="MSTG-NET-2",
            affected_component="API Communication Layer"
        )

        self.scan_result.add_finding(finding)
        self.logger.info("DAST tests completed")

    def _run_frida_tests(self):
        """Execute instrumentation tests with Frida"""
        self.logger.info("Running Frida instrumentation tests...")

        # Sample findings
        finding = Finding(
            id="FINDING-003",
            title="Root Detection Bypass",
            description="Root detection mechanism can be bypassed",
            severity=Severity.MEDIUM,
            cvss=self._create_cvss_score(5.5),
            cwe=["CWE-656"],
            owasp_category="A06:2021 - Vulnerable and Outdated Components",
            test_name="Root Detection Bypass",
            module="frida",
            mastg_category="MASTG-RESILIENCE-1",
            masvs_category="MSTG-RESILIENCE-1",
            affected_component="Security Check Method"
        )

        self.scan_result.add_finding(finding)
        self.logger.info("Frida tests completed")

    def _create_cvss_score(self, score: float) -> Dict:
        """Create CVSS score object"""
        vectors = {
            9.8: "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H",
            7.5: "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N",
            5.5: "CVSS:3.1/AV:L/AC:L/PR:L/UI:N/S:U/C:H/I:N/A:N",
        }
        return {
            "score": score,
            "vector": vectors.get(score, f"CVSS:3.1/AV:L/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N"),
        }

    def _map_findings_to_masvs(self):
        """Map findings to MASVS requirements"""
        self.logger.info("Mapping findings to MASVS requirements...")

        for finding in self.scan_result.findings:
            # Find related MASVS requirements
            masvs_reqs = MAVSMapping.map_mastg_to_masvs(finding.mastg_category)

            if masvs_reqs:
                finding.masvs_category = masvs_reqs[0].id

        # Calculate MASVS compliance
        self._calculate_masvs_compliance()

    def _calculate_masvs_compliance(self):
        """Calculate MASVS compliance levels"""
        if not self.scan_result:
            return

        # Get all MASVS requirements by level
        l1_requirements = MAVSMapping.get_requirements_by_level(MAVSLevel.L1)
        l2_requirements = MAVSMapping.get_requirements_by_level(MAVSLevel.L2)
        r_requirements = MAVSMapping.get_requirements_by_level(MAVSLevel.R)

        # Build set of failed requirements from findings
        failed_requirements = set()
        for finding in self.scan_result.findings:
            if finding.masvs_category:
                failed_requirements.add(finding.masvs_category)

        # Calculate compliance for L1
        l1_passed = sum(1 for req in l1_requirements if req.id not in failed_requirements)
        l1_coverage = l1_passed / len(l1_requirements) * 100 if l1_requirements else 0

        # Calculate compliance for L2
        l2_passed = sum(1 for req in l2_requirements if req.id not in failed_requirements)
        l2_coverage = l2_passed / len(l2_requirements) * 100 if l2_requirements else 0

        # Calculate compliance for R
        r_passed = sum(1 for req in r_requirements if req.id not in failed_requirements)
        r_coverage = r_passed / len(r_requirements) * 100 if r_requirements else 0

        # Store compliance in scan result
        self.scan_result.masvs_compliance = {
            "L1": {
                "coverage": round(l1_coverage, 2),
                "passed": l1_passed,
                "total": len(l1_requirements)
            },
            "L2": {
                "coverage": round(l2_coverage, 2),
                "passed": l2_passed,
                "total": len(l2_requirements)
            },
            "R": {
                "coverage": round(r_coverage, 2),
                "passed": r_passed,
                "total": len(r_requirements)
            }
        }

        self.logger.info(f"MASVS Compliance - L1: {l1_coverage:.1f}% | L2: {l2_coverage:.1f}% | R: {r_coverage:.1f}%")

    def generate_report(self, format: str = "json") -> str:
        """
        Generate report in specified format.

        Args:
            format: Report format (json, pdf, docx, markdown)

        Returns:
            str: Report content or filepath
        """
        if not self.scan_result:
            raise RuntimeError("No scan result available")

        self.logger.info(f"Generating {format} report...")

        if format == "json":
            return self.scan_result.to_json()
        elif format == "pdf":
            return self._generate_pdf_report()
        elif format == "docx":
            return self._generate_docx_report()
        elif format == "markdown":
            return self._generate_markdown_report()
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_pdf_report(self) -> str:
        """Generate PDF report using reportlab"""
        self.logger.info("Generating PDF report...")

        if not REPORTLAB_AVAILABLE:
            self.logger.warning("reportlab not available, falling back to markdown format")
            return self._generate_markdown_report()

        try:
            filename = f"mobscan_report_{self.scan_result.scan_id}.pdf"
            filepath = Path(filename)

            # Create PDF document
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            story = []
            styles = getSampleStyleSheet()

            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=30,
                alignment=1  # Center
            )
            story.append(Paragraph("Security Assessment Report", title_style))
            story.append(Spacer(1, 0.3*inch))

            # Executive Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            summary_data = [
                ["Application Name", self.scan_result.app_info.app_name],
                ["Package Name", self.scan_result.app_info.package_name],
                ["Platform", self.scan_result.app_info.platform.upper()],
                ["Scan Date", self.scan_result.started_at.isoformat()],
                ["Risk Score", f"{self.scan_result.risk_metrics.risk_score}/10"],
            ]
            summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8eef7')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ]))
            story.append(summary_table)
            story.append(Spacer(1, 0.3*inch))

            # Findings Summary
            story.append(Paragraph("Findings Summary", styles['Heading2']))
            findings_data = [
                ["Severity", "Count"],
                ["Critical", str(self.scan_result.risk_metrics.critical_count)],
                ["High", str(self.scan_result.risk_metrics.high_count)],
                ["Medium", str(self.scan_result.risk_metrics.medium_count)],
                ["Low", str(self.scan_result.risk_metrics.low_count)],
                ["Info", str(self.scan_result.risk_metrics.info_count)],
            ]
            findings_table = Table(findings_data, colWidths=[2*inch, 2*inch])
            findings_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ]))
            story.append(findings_table)
            story.append(Spacer(1, 0.3*inch))

            # Detailed Findings
            if self.scan_result.findings:
                story.append(PageBreak())
                story.append(Paragraph("Detailed Findings", styles['Heading2']))

                for finding in self.scan_result.findings:
                    # Finding title
                    finding_title = f"{finding.title} ({finding.severity.value})"
                    story.append(Paragraph(finding_title, styles['Heading3']))

                    # Finding details
                    details_data = [
                        ["CVSS Score", str(finding.cvss.get('score', 'N/A'))],
                        ["Component", finding.affected_component],
                        ["MASTG Reference", finding.mastg_category],
                    ]
                    details_table = Table(details_data, colWidths=[1.5*inch, 4.5*inch])
                    details_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ]))
                    story.append(details_table)

                    # Finding description
                    story.append(Paragraph(finding.description, styles['Normal']))
                    story.append(Spacer(1, 0.2*inch))

            # Build PDF
            doc.build(story)
            self.logger.info(f"PDF report generated: {filename}")
            return filename

        except Exception as e:
            self.logger.error(f"Error generating PDF report: {e}")
            return self._generate_markdown_report()

    def _generate_docx_report(self) -> str:
        """Generate DOCX report using python-docx"""
        self.logger.info("Generating DOCX report...")

        if not PYTHON_DOCX_AVAILABLE:
            self.logger.warning("python-docx not available, falling back to markdown format")
            return self._generate_markdown_report()

        try:
            filename = f"mobscan_report_{self.scan_result.scan_id}.docx"

            # Create document
            doc = Document()

            # Title
            title = doc.add_heading("Security Assessment Report", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Executive Summary
            doc.add_heading("Executive Summary", level=2)
            summary_table = doc.add_table(rows=6, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            summary_data = [
                ["Application Name", self.scan_result.app_info.app_name],
                ["Package Name", self.scan_result.app_info.package_name],
                ["Platform", self.scan_result.app_info.platform.upper()],
                ["Scan Date", self.scan_result.started_at.isoformat()],
                ["Risk Score", f"{self.scan_result.risk_metrics.risk_score}/10"],
            ]
            for i, (key, value) in enumerate(summary_data):
                summary_table.rows[i].cells[0].text = key
                summary_table.rows[i].cells[1].text = value

            doc.add_paragraph()

            # Findings Summary
            doc.add_heading("Findings Summary", level=2)
            findings_table = doc.add_table(rows=7, cols=2)
            findings_table.style = 'Light Grid Accent 1'
            findings_table.rows[0].cells[0].text = "Severity"
            findings_table.rows[0].cells[1].text = "Count"
            findings_table.rows[1].cells[0].text = "Critical"
            findings_table.rows[1].cells[1].text = str(self.scan_result.risk_metrics.critical_count)
            findings_table.rows[2].cells[0].text = "High"
            findings_table.rows[2].cells[1].text = str(self.scan_result.risk_metrics.high_count)
            findings_table.rows[3].cells[0].text = "Medium"
            findings_table.rows[3].cells[1].text = str(self.scan_result.risk_metrics.medium_count)
            findings_table.rows[4].cells[0].text = "Low"
            findings_table.rows[4].cells[1].text = str(self.scan_result.risk_metrics.low_count)
            findings_table.rows[5].cells[0].text = "Info"
            findings_table.rows[5].cells[1].text = str(self.scan_result.risk_metrics.info_count)

            doc.add_page_break()

            # Detailed Findings
            if self.scan_result.findings:
                doc.add_heading("Detailed Findings", level=2)

                for finding in self.scan_result.findings:
                    # Finding title
                    doc.add_heading(f"{finding.title} ({finding.severity.value})", level=3)

                    # Finding details table
                    details_table = doc.add_table(rows=4, cols=2)
                    details_table.style = 'Light Grid Accent 1'
                    details_table.rows[0].cells[0].text = "CVSS Score"
                    details_table.rows[0].cells[1].text = str(finding.cvss.get('score', 'N/A'))
                    details_table.rows[1].cells[0].text = "Component"
                    details_table.rows[1].cells[1].text = finding.affected_component
                    details_table.rows[2].cells[0].text = "MASTG Reference"
                    details_table.rows[2].cells[1].text = finding.mastg_category
                    details_table.rows[3].cells[0].text = "CWE"
                    details_table.rows[3].cells[1].text = ", ".join(finding.cwe) if finding.cwe else "N/A"

                    # Finding description
                    doc.add_paragraph(finding.description)
                    doc.add_paragraph()

            # Save document
            doc.save(filename)
            self.logger.info(f"DOCX report generated: {filename}")
            return filename

        except Exception as e:
            self.logger.error(f"Error generating DOCX report: {e}")
            return self._generate_markdown_report()

    def _generate_markdown_report(self) -> str:
        """Generate Markdown report"""
        self.logger.info("Generating Markdown report...")

        md = f"""# Security Assessment Report

## Executive Summary

**App**: {self.scan_result.app_info.app_name}
**Package**: {self.scan_result.app_info.package_name}
**Platform**: {self.scan_result.app_info.platform.upper()}
**Scan Date**: {self.scan_result.started_at.isoformat()}
**Risk Score**: {self.scan_result.risk_metrics.risk_score}/10

### Findings Summary

- **Critical**: {self.scan_result.risk_metrics.critical_count}
- **High**: {self.scan_result.risk_metrics.high_count}
- **Medium**: {self.scan_result.risk_metrics.medium_count}
- **Low**: {self.scan_result.risk_metrics.low_count}
- **Info**: {self.scan_result.risk_metrics.info_count}

## Detailed Findings

"""

        for finding in self.scan_result.findings:
            md += f"""### {finding.title}

**Severity**: {finding.severity.value}
**CVSS Score**: {finding.cvss.get('score', 'N/A')}
**Component**: {finding.affected_component}
**MASTG Reference**: {finding.mastg_category}

{finding.description}

"""

        return md

    def save_scan_result(self, filepath: str):
        """Save scan result to file"""
        if not self.scan_result:
            raise RuntimeError("No scan result to save")

        self.scan_result.export_to_json(filepath)
        self.logger.info(f"Scan result saved to {filepath}")

    def get_scan_statistics(self) -> Dict[str, Any]:
        """Get comprehensive scan statistics"""
        if not self.scan_result:
            return {}

        return {
            "scan_id": self.scan_result.scan_id,
            "duration_seconds": self.scan_result.duration_seconds,
            "total_findings": len(self.scan_result.findings),
            "risk_metrics": self.scan_result.risk_metrics.to_dict(),
            "test_coverage": self.scan_result.test_coverage.to_dict(),
        }
